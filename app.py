# ==============================================================================
# Multi-Sensor Data Explorer
#
# This script creates a Streamlit web application for visualising and analysing
# time-series data from multiple sensors. It features data loading from a
# database, interactive plotting, statistical analysis, and the ability to
# upload new data files. It can also generate formatted Word documents of the
# plots for reporting purposes.
# ==============================================================================

# --- Core Python Libraries ---
# These modules are part of Python's standard library and provide essential tools.
import os
import shutil
import time
import logging
import math
import io
from datetime import datetime, date
from typing import List, Tuple, Optional, Dict, Any

# --- Third-Party Libraries ---
# These are external packages that need to be installed (e.g., via pip).
# They provide the powerful functionality for data handling, plotting, etc.
import apsw
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.styles.styles import Styles
from docx.oxml.ns import qn

# --- Local Module Imports ---
# Attempt to import custom modules for database and data handling.
# If they are not found, the script will fall back to using mock classes
# for demonstration and standalone execution.
try:
    from database_handler import DatabaseHandler, DB_PATH
    from data_loader import DataLoader
except ImportError:
    # --- Mock Classes for Standalone Execution ---
    # These classes simulate the behaviour of the real DatabaseHandler and
    # DataLoader. This allows the Streamlit application to run and display
    # sample data even without the actual database and data loading logic.
    class MockDBHandler:
        """A mock database handler that generates sample data."""
        def load_all_data(self) -> pd.DataFrame:
            st.warning("Using mock data. `database_handler` not found.")
            try:
                import numpy as np
            except ImportError:
                st.error("Numpy is required for mock data generation.")
                return pd.DataFrame()
            # Generate a date range for our sample data.
            dates: pd.Series = pd.to_datetime(pd.date_range(start="2024-01-01", end="2025-03-31", freq='h'))
            data: Dict[str, Any] = {'datetime': dates}
            # Create several channels of noisy, trending data.
            for i in range(8):
                noise = (np.random.randn(len(dates)) * 0.1 * (i + 1)).cumsum()
                trend = np.linspace(start=0, stop=i*0.1, num=len(dates))
                data[f'CH{i+1}_profile'] = 0.5 + trend + noise
            return pd.DataFrame(data)
        
        def get_latest_datetime(self) -> pd.Timestamp:
            return pd.to_datetime("2025-03-31")
        
        def save_to_db(self, df: pd.DataFrame) -> None:
            pass

    class MockDataLoader:
        """A mock data loader that simulates file processing."""
        def __init__(self, uploaded_file: Any):
            pass
        def load_and_clean(self) -> pd.DataFrame:
            return pd.DataFrame()
        def compute_channel_differences(self, df: pd.DataFrame, factors: List[float]) -> pd.DataFrame:
            return df

    # Assign the mock classes to be used by the rest of the script.
    DatabaseHandler = MockDBHandler # type: ignore
    DataLoader = MockDataLoader # type: ignore
    DB_PATH: str = "mock.db"


# --- Global Configuration ---

# Set up basic logging to output informational messages to the console.
logging.basicConfig(level=logging.INFO)
logger: logging.Logger = logging.getLogger(__name__)

# Define a fixed colour palette for the plots to ensure consistent and
# visually distinct traces for each sensor channel.
PALETTE: List[str] = [
    "#005eff", "#5500FF", '#ff7c43', '#ffa600',
    "#00FF00", "#FF0000", "#ff00bf", "#00D9FF",
    "#333333", "#007044",
]


# --- Plotting Utility Functions ---

def get_plotly_layout() -> Dict[str, Any]:
    """
    Defines a base layout template for all Plotly figures.
    This ensures a consistent, professional look and feel (e.g., white
    background, clear fonts) across all visualisations in the application.
    """
    font_colour: str = '#111111'
    return {
        'template': 'plotly',
        'plot_bgcolor': 'white',
        'paper_bgcolor': 'white',
        'font': {'color': font_colour},
        'xaxis': {'tickfont': {'color': font_colour}, 'title': {'font': {'color': font_colour}}},
        'yaxis': {'tickfont': {'color': font_colour}, 'title': {'font': {'color': font_colour}}},
        'legend': {'font': {'color': font_colour}, 'title': {'font': {'color': font_colour}}},
        'hovermode': 'x unified'
    }


def calculate_flexible_y_range(series: pd.Series, padding_factor: float = 0.2) -> Tuple[float, float, Optional[float]]:
    """
    Calculates a dynamic, visually appealing y-axis range for a plot.
    Instead of letting the plot library auto-scale, this function finds the
    min/max of the data, adds some padding, and then calculates "nice" rounded
    tick intervals to make the graph easy to read.
    """
    data_min: float = series.min()
    data_max: float = series.max()

    # Handle cases with no variation in data.
    if data_min == data_max:
        padding: float = abs(data_min * padding_factor) if data_min != 0 else 0.1
        return data_min - padding, data_max + padding, None

    span: float = data_max - data_min
    if span < 1e-9: # Handle very small spans.
        padding = abs(data_min * 0.1) if data_min != 0 else 0.1
        return data_min - padding, data_max + padding, padding / 2

    # Calculate padding and a "nice" step for the axis ticks.
    padding = span * padding_factor
    padded_min: float = data_min - padding
    padded_max: float = data_max + padding
    padded_span: float = padded_max - padded_min

    target_ticks: int = 8
    raw_step: float = padded_span / target_ticks
    power: float = 10.0 ** math.floor(math.log10(raw_step))
    normalised_step: float = raw_step / power
    multiples: List[float] = [1, 2, 2.5, 5, 10]
    best_multiple: float = min(multiples, key=lambda x: abs(x - normalised_step))
    final_dtick: float = best_multiple * power

    # Calculate the final rounded min and max for the axis range.
    final_min: float = math.floor(padded_min / final_dtick) * final_dtick
    final_max: float = math.ceil(padded_max / final_dtick) * final_dtick

    return final_min, final_max, final_dtick


def create_line_figure(
    df: pd.DataFrame,
    x_col: str,
    y_col: str,
    y_label: str,
    colour: str,
    title_text: str,
    y_range: Tuple[float, float],
    y_dtick: Optional[float],
    tick_format: str = ".3f",
    font_sizes: Optional[Dict[str, int]] = None
) -> go.Figure:
    """
    Builds a self-contained, fully styled Plotly line figure for an individual channel.
    This function uses `plotly.graph_objects` (go) for its core plotting command,
    as it provides a more robust and direct way to handle data types like dates,
    preventing formatting issues when exporting to static images.
    """
    fig: go.Figure = go.Figure()

    # Add the line trace (the actual data) to the figure.
    fig.add_trace(go.Scatter(
        x=df[x_col],
        y=df[y_col],
        mode='lines',
        name=y_col,
        line=dict(color=colour),
        showlegend=True
    ))

    # Configure the y-axis with the calculated range and gridlines.
    fig.update_yaxes(
        title_text=y_label,
        range=y_range,
        tickformat=tick_format,
        dtick=y_dtick,
        showgrid=True,
        gridcolor='lightgrey',
        gridwidth=1
    )
    # Configure the x-axis, providing an explicit tick format for dates.
    fig.update_xaxes(
        title_text='Date & Time',
        showgrid=True,
        gridcolor='lightgrey',
        gridwidth=1,
        tickformat='%Y-%m-%d'  # Explicit format for robust date rendering.
    )

    # Apply the base layout and then customise it with specific titles and fonts.
    layout_config: Dict[str, Any] = get_plotly_layout()
    defaults: Dict[str, int] = {'base': 14, 'legend': 16, 'axis_title': 16, 'axis_tick': 14, 'title': 18}
    fonts: Dict[str, int] = {**defaults, **(font_sizes or {})}
    
    layout_config['font']['size'] = fonts['base']
    layout_config['legend']['font']['size'] = fonts['legend']
    layout_config['xaxis']['tickfont']['size'] = fonts['axis_tick']
    layout_config['yaxis']['tickfont']['size'] = fonts['axis_tick']
    layout_config['legend'].update(orientation='v', x=1.02, xanchor='left', y=1, yanchor='top')
    layout_config['title'] = {'text': title_text, 'x': 0.5, 'xanchor': 'center', 'yanchor': 'top', 'y': 0.95, 'font': {'size': fonts['title'], 'color': '#111111'}}
    layout_config['margin'] = {'t': 100, 'r': 300}

    fig.update_layout(layout_config)
    
    return fig

# --- Word Document Generation ---
# def set_aptos_font(doc: Document) -> None:
#     """
#     Explicitly finds key document styles by name and sets their font to Aptos.
#     This is a more robust method that directly targets the styles used for
#     headings, paragraphs, and captions to ensure the font is changed correctly.
#     """
#     # A list of the specific style names we are using in the document.
#     style_names_to_change: List[str] = [
#         'Normal',
#         'Title',
#         'Heading 1',
#         'Heading 2',
#         'Heading 3',
#         'List Paragraph',
#         'Para',
#         'Body Text',
#         'Caption'
#     ]

#     for style_name in style_names_to_change:
#         try:
#             # The correct syntax is to access styles using square brackets,
#             # like accessing a dictionary key.
#             style = doc.styles[style_name]
            
#             if hasattr(style, 'font'):
#                 font = style.font
#                 font.name = 'Aptos'

#         except KeyError:
#             # If a style (e.g., 'Title') doesn't exist in the base document,
#             # this prevents a crash and logs a minor warning.
#             logger.warning(f"Style '{style_name}' not found in document. Skipping font change.")
#             pass
# You must have this import at the top of your file
from docx.oxml.ns import qn

def set_aptos_font(doc: Document) -> None:
    """
    Applies the Aptos font using a hybrid approach to ensure all styles,
    including theme-based heading styles, are correctly updated.
    """
    # Part 1: Apply the simple, high-level fix for non-theme styles.
    # This works for styles like 'Normal' and 'Caption'.
    simple_styles = ['Normal', 'Caption']
    for style_name in simple_styles:
        try:
            style = doc.styles[style_name]
            if hasattr(style, 'font'):
                style.font.name = 'Aptos'
        except KeyError:
            logger.warning(f"Style '{style_name}' not found; skipping.")

    # Part 2: Apply a forceful, low-level fix for theme-based heading styles.
    # This directly manipulates the style's underlying XML to override the
    # theme font (e.g., 'Calibri (Headings)') with a direct font ('Aptos').
    heading_styles = ['Title', 'Heading 1', 'Heading 2', 'Heading 3']
    for style_name in heading_styles:
        try:
            style = doc.styles[style_name]
            # Access the style's underlying XML properties.
            rpr = style.element.get_or_add_rPr()
            # Clear any existing theme font reference.
            rpr.rFonts.attrib.clear()
            # Set the font directly for all character types.
            rpr.rFonts.set(qn('w:ascii'), 'Aptos')
            rpr.rFonts.set(qn('w:hAnsi'), 'Aptos')
        except KeyError:
            logger.warning(f"Style '{style_name}' not found; skipping.")


def generate_word_docs(
    combined_fig: go.Figure,
    individual_figs: List[go.Figure],
    stats_content: List[Tuple[str, Dict[str, float], go.Figure]]
) -> Tuple[io.BytesIO, io.BytesIO]:
    """
    Generates two Word documents in memory containing all plots and statistics.

    This function takes the generated Plotly figure objects, performs a final
    data type conversion on their date axes to ensure compatibility with the
    backend image renderer, converts them to PNG images, and inserts them into
    .docx files with the Aptos font applied to all text.
    """
    # --- Document 1: Sensor Graphs ---
    # Initialise a new, blank Word document object.
    graphs_doc: Document = Document()
    set_aptos_font(graphs_doc)  # Apply the universal font style.
    
    graphs_doc.add_heading('Appendix A', level=0)
    graphs_doc.add_heading('Sensor Graphs', level=1)
    graphs_doc.add_page_break()
    
    # --- Targeted Fix for Combined Graph ---
    # This loop surgically modifies the figure object just before saving.
    # It converts the x-axis data from NumPy's datetime64 to Python's native
    # datetime object, which the backend image renderer handles perfectly.
    # This leaves the original figure used by the web app untouched.
    for trace in combined_fig.data:
        trace.x = pd.to_datetime(trace.x).to_pydatetime()

    # Convert the modified figure to a PNG image represented as bytes.
    img_bytes: bytes = combined_fig.to_image(format="png", width=800, height=500)
    graphs_doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
    graphs_doc.add_paragraph('Figure 1: Combined Normalised Graph', style='Caption')

    # Now, repeat the process for each of the individual graphs.
    for idx, fig in enumerate(individual_figs):
        channel_name: str = fig.layout.title.text.split(':')[0]
        
        # Apply the same targeted data conversion fix to each individual figure.
        for trace in fig.data:
            trace.x = pd.to_datetime(trace.x).to_pydatetime()

        img_bytes = fig.to_image(format="png", width=800, height=500)
        graphs_doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
        caption_text: str = f"Figure {idx + 2}: {channel_name}"
        graphs_doc.add_paragraph(caption_text, style='Caption')

    # Save the completed document to an in-memory binary stream (a buffer).
    graphs_io: io.BytesIO = io.BytesIO()
    graphs_doc.save(graphs_io)
    # Rewind the buffer to the beginning so it can be read by Streamlit.
    graphs_io.seek(0)

    # --- Document 2: Channel Statistics ---
    stats_doc: Document = Document()
    set_aptos_font(stats_doc)  # Apply the universal font style.
    
    stats_doc.add_heading('Channel Statistics', 0)
    for channel, stats, hist_fig in stats_content:
        stats_doc.add_heading(f"Statistics for {channel}", level=1)
        stats_doc.add_paragraph(f"Mean: {stats['mean']:.3f}")
        stats_doc.add_paragraph(f"Standard Deviation: {stats['std']:.3f}")
        stats_doc.add_paragraph(f"Minimum: {stats['min']:.3f}")
        stats_doc.add_paragraph(f"Maximum: {stats['max']:.3f}")
        
        stats_doc.add_heading(f"Distribution of {channel}", level=2)
        img_bytes = hist_fig.to_image(format="png", width=700, height=400)
        stats_doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
        stats_doc.add_page_break()

    stats_io: io.BytesIO = io.BytesIO()
    stats_doc.save(stats_io)
    stats_io.seek(0)

    return graphs_io, stats_io


# --- Streamlit Rendering Functions ---

def render_individual_graphs(
    df_filtered: pd.DataFrame,
    selected_channels: List[str],
    start_date: date,
    end_date: date
) -> List[go.Figure]:
    """
    Renders a separate graph for each selected channel onto the Streamlit page.
    """
    st.subheader("📉 Individual Channel Graphs")
    figures: List[go.Figure] = []
    for idx, channel in enumerate(selected_channels):
        st.markdown(f"### {channel}")
        if channel not in df_filtered.columns:
            st.warning(f"⚠️ No data for {channel}")
            continue

        # Prepare the data for this specific channel.
        df_ch: pd.DataFrame = df_filtered[['datetime', channel]].dropna()
        df_ch[channel] = pd.to_numeric(df_ch[channel], errors='coerce')
        df_ch.dropna(subset=[channel], inplace=True)
        
        if df_ch.empty or df_ch[channel].isnull().all():
            st.warning(f"⚠️ No valid numeric data for {channel}")
            continue

        # Generate the figure using our utility function.
        y_min, y_max, y_dtick = calculate_flexible_y_range(df_ch[channel])
        y_range: Tuple[float, float] = (y_min, y_max)
        colour: str = PALETTE[idx % len(PALETTE)]
        title_text: str = f"{channel}: {start_date} to {end_date}"
        y_label: str = 'Crack width change [mm]'

        fig: go.Figure = create_line_figure(
            df=df_ch, x_col='datetime', y_col=channel, y_label=y_label,
            colour=colour, title_text=title_text, y_range=y_range, y_dtick=y_dtick
        )
        figures.append(fig)
        st.plotly_chart(fig, use_container_width=True)
    return figures


def render_combined_normalised_graph(
    df: pd.DataFrame,
    all_channel_cols: List[str],
    start_date_default: date,
    end_date_default: date
) -> go.Figure:
    """
    Renders a single graph containing all selected channels, plotted together.
    """
    st.subheader("📈 Combined Normalised Graph")
    if df.empty:
        st.warning("📪 No records for selected range")
        return go.Figure()

    # Create an empty figure to which we will add traces in a loop.
    fig: go.Figure = go.Figure()
    for idx, channel in enumerate(all_channel_cols):
        fig.add_trace(go.Scatter(
            x=df['datetime'],
            y=df[channel],
            mode='lines',
            name=channel,
            line=dict(color=PALETTE[idx % len(PALETTE)])
        ))

    # Define plot-specific labels and ranges.
    y_label: str = 'Crack width change [mm]'
    title_text: str = f"Sensor Data from {start_date_default} to {end_date_default}"
    y_range: Tuple[float, float] = (-0.8, 1.6)

    # Configure axes with an explicit tick format for dates.
    fig.update_yaxes(
        title_text=y_label,
        range=y_range,
        tickformat=".3f",
        dtick=0.2,
        showgrid=True,
        gridcolor='lightgrey',
        gridwidth=1
    )
    fig.update_xaxes(
        title_text='Date & Time',
        showgrid=True,
        gridcolor='lightgrey',
        gridwidth=1,
        tickformat='%Y-%m-%d'
    )
    
    # Apply the base layout and custom styling.
    layout_config: Dict[str, Any] = get_plotly_layout()
    defaults: Dict[str, int] = {'base': 14, 'legend': 16, 'axis_title': 16, 'axis_tick': 14, 'title': 18}
    layout_config['font']['size'] = defaults['base']
    layout_config['legend']['font']['size'] = defaults['legend']
    layout_config['xaxis']['tickfont']['size'] = defaults['axis_tick']
    layout_config['yaxis']['tickfont']['size'] = defaults['axis_tick']
    layout_config['legend'].update(orientation='v', x=1.02, xanchor='left', y=1, yanchor='top', title_text='Sensor')
    layout_config['title'] = dict(
        text=title_text, x=0.5, xanchor='center', yanchor='top', y=0.95,
        font=dict(size=defaults['title'], color='#111111')
    )
    layout_config['margin'] = dict(t=60, r=300)

    fig.update_layout(layout_config)
    st.plotly_chart(fig, use_container_width=True)
    return fig


def render_statistics(
    df: pd.DataFrame,
    columns: List[str]
) -> List[Tuple[str, Dict[str, float], go.Figure]]:
    """
    Renders statistical summaries (mean, std dev, etc.) and a histogram
    for each selected channel.
    """
    st.subheader("📊 Channel Statistics")
    
    all_stats_content: List[Tuple[str, Dict[str, float], go.Figure]] = []
    
    if not columns:
        st.warning("No channels selected to generate statistics.")
        return []

    for sel in columns:
        st.markdown(f"#### Statistics for {sel}")
        vals: pd.Series = pd.to_numeric(df[sel], errors='coerce').dropna()
        if vals.empty:
            st.warning(f"⚠️ No numeric data for {sel}.")
            continue
        
        # Calculate summary statistics.
        mean: float = vals.mean()
        std: float = vals.std()
        vmin: float = vals.min()
        vmax: float = vals.max()
        
        stats_dict: Dict[str, float] = {'mean': mean, 'std': std, 'min': vmin, 'max': vmax}

        # Display stats and histogram in columns for a compact layout.
        col1, col2 = st.columns([1, 2])
        with col1:
            st.markdown(f"**Mean:** {mean:.3f}")
            st.markdown(f"**Std Dev:** {std:.3f}")
            st.markdown(f"**Min:** {vmin:.3f}")
            st.markdown(f"**Max:** {vmax:.3f}")

        hist: go.Figure = go.Figure(data=[go.Histogram(x=vals, marker_color=PALETTE[columns.index(sel) % len(PALETTE)])])
        
        # Configure histogram axes and layout.
        hist.update_xaxes(title_text=sel)
        hist.update_yaxes(title_text="Count")
        
        hist.update_layout(
            **get_plotly_layout(),
            title=f"Distribution of {sel}",
            margin=dict(t=50, l=10, r=10, b=10)
        )

        with col2:
            st.plotly_chart(hist, use_container_width=True)
        
        all_stats_content.append((sel, stats_dict, hist))

    return all_stats_content


def backup_database(
    db_path: str = DB_PATH,
    backup_dir: str = "backups"
) -> Optional[str]:
    """
    Creates a timestamped backup of the database file before new data is added.
    """
    if not os.path.exists(db_path) or db_path == "mock.db":
        logger.warning(f"DB missing or using mock at: {db_path}. Skipping backup.")
        return None
    try:
        conn: apsw.Connection = apsw.Connection(db_path)
        cur: apsw.Cursor = conn.cursor()
        # Check if there are any tables to back up.
        tables: List[str] = [r[0] for r in cur.execute("SELECT name FROM sqlite_master WHERE type='table';")]
        conn.close()
        if not tables:
            logger.info("No tables; skip backup.")
            return None
        # Create the backup file.
        ts: str = datetime.now().strftime('%Y%m%d_%H%M%S')
        script_dir: str = os.path.dirname(os.path.abspath(__file__)) if '__file__' in locals() else os.getcwd()
        out: str = os.path.join(script_dir, backup_dir)
        os.makedirs(out, exist_ok=True)
        dst: str = os.path.join(out, f"backup_{ts}.db")
        shutil.copy2(db_path, dst)
        logger.info(f"Backed up to {dst}")
        return dst
    except Exception as e:
        logger.error(f"Database backup failed: {e}")
        st.error(f"Database backup failed: {e}")
        return None

# ==============================================================================
# --- Streamlit Application Main Logic ---
# ==============================================================================

# --- Page Title and Main Action Button ---
title_col, button_col = st.columns([3, 1])
with title_col:
    st.title("📊 Multi-Sensor Data Explorer")

# --- Session State Initialisation ---
# `st.session_state` is a dictionary-like object that persists across reruns,
# allowing us to store state, such as generated files or user selections.
if 'docs_ready' not in st.session_state:
    st.session_state.docs_ready = False
if 'graphs_file_data' not in st.session_state:
    st.session_state.graphs_file_data = None
if 'stats_file_data' not in st.session_state:
    st.session_state.stats_file_data = None
    
st.sidebar.radio("Theme", ["Light", "Dark"], index=0, disabled=False) 

# --- Handle Print Button Click & Persistent Downloads ---
if button_col.button("🖨️ Print Plots"):
    # Check if the figures have been generated and stored in the session state.
    if 'combined_fig' in st.session_state and st.session_state.combined_fig:
        with st.spinner("Generating Word documents... This may take a moment."):
            graphs_file, stats_file = generate_word_docs(
                st.session_state.combined_fig,
                st.session_state.individual_figs,
                st.session_state.stats_content
            )
            # Store the generated documents in the session state.
            st.session_state.graphs_file_data = graphs_file
            st.session_state.stats_file_data = stats_file
            st.session_state.docs_ready = True
    else:
        st.warning("No plots available to print. Please ensure data is loaded and filters are applied.")

# This block *reads* from the session state. If `docs_ready` is True, it will
# show the download buttons, making them persist after the first click.
if st.session_state.get('docs_ready'):
    st.success("✅ Documents are ready for download!")
    
    dl_col1, dl_col2 = st.columns(2)
    with dl_col1:
        st.download_button(
            label="📄 Download Graphs Document",
            data=st.session_state.graphs_file_data,
            file_name="Sensor_Graphs.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with dl_col2:
        st.download_button(
            label="📊 Download Statistics Document",
            data=st.session_state.stats_file_data,
            file_name="Sensor_Statistics.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- Main Data Loading and Display Logic ---
db: DatabaseHandler = DatabaseHandler()
df_all: Optional[pd.DataFrame] = db.load_all_data()

if df_all is not None and not df_all.empty:
    # This is a critical step: ensure the datetime column is the correct data type.
    df_all['datetime'] = pd.to_datetime(df_all['datetime'])
    df_all = df_all.sort_values("datetime")
    all_channel_cols: List[str] = [col for col in df_all.columns if col.endswith("profile")]

    # Initialise filter widgets with values from the session state, or defaults.
    if "start_date" not in st.session_state:
        st.session_state["start_date"] = df_all["datetime"].min().date()
    if "end_date" not in st.session_state:
        st.session_state["end_date"] = df_all["datetime"].max().date()
    if "selected_channels" not in st.session_state:
        st.session_state["selected_channels"] = all_channel_cols[:]

    # Create a form for the filter controls to batch user inputs.
    with st.form("filter_form"):
        start_date_input: date = st.date_input("Start Date", st.session_state["start_date"])
        end_date_input: date = st.date_input("End Date", st.session_state["end_date"])
        selected_channels_input: List[str] = st.multiselect(
            "Select Channels",
            options=all_channel_cols,
            default=st.session_state["selected_channels"],
        )
        if st.form_submit_button("Apply Filters"):
            # When the form is submitted, update the session state with the new values.
            st.session_state["start_date"] = start_date_input
            st.session_state["end_date"] = end_date_input
            st.session_state["selected_channels"] = selected_channels_input
            
            # Invalidate any previously generated documents, as the data has changed.
            st.session_state.docs_ready = False
            st.session_state.graphs_file_data = None
            st.session_state.stats_file_data = None
            
            st.rerun()

    # Filter the main DataFrame based on the selections stored in the session state.
    if not st.session_state["selected_channels"]:
        st.warning("⚠️ No channels selected. Please select at least one channel.")
    else:
        df_filtered: pd.DataFrame = df_all[
            (df_all['datetime'].dt.date >= st.session_state["start_date"]) &
            (df_all['datetime'].dt.date <= st.session_state["end_date"])
        ]
        
        if df_filtered.empty:
            st.warning("📪 No records for selected range")
        else:
            selected_data: pd.DataFrame = df_filtered[['datetime'] + st.session_state["selected_channels"]]
            
            # Call the rendering functions to display the plots and stats.
            # Store the returned figure objects in the session state so the
            # "Print Plots" button can access them.
            st.session_state.combined_fig = render_combined_normalised_graph(
                selected_data,
                st.session_state["selected_channels"],
                st.session_state["start_date"],
                st.session_state["end_date"]
            )
            st.session_state.individual_figs = render_individual_graphs(
                selected_data,
                st.session_state["selected_channels"],
                st.session_state["start_date"],
                st.session_state["end_date"],
            )
            st.session_state.stats_content = render_statistics(selected_data, st.session_state["selected_channels"])
else:
    st.error("Could not load any data. Please check the database connection.")

# --- File Uploader Section ---
st.markdown("---")
st.subheader("📂 Upload New Sensor File")

uploaded_file: Optional[Any] = st.file_uploader(
    "Upload Excel or CSV File",
    type=["csv", "xlsx", "xls"],
    key="uploader"
)
if uploaded_file is not None:
    # Create a "fingerprint" of the file to avoid reprocessing the same file.
    fingerprint: str = f"{uploaded_file.name}_{uploaded_file.size}"
    if st.session_state.get("last_upload") == fingerprint:
        st.info("⚠️ This file was already processed.")
    else:
        # --- Data Ingestion Pipeline ---
        st.write("▶️ Step 1: Backing up database…")
        t0: float = time.time()
        backup_database()
        st.write(f"✅ Backup done in {time.time() - t0:.2f}s")

        st.write("▶️ Step 2: Loading & cleaning file…")
        t1: float = time.time()
        loader: DataLoader = DataLoader(uploaded_file)
        df_new: pd.DataFrame = loader.load_and_clean()
        st.write(f"✅ load_and_clean done in {time.time() - t1:.2f}s (shape={df_new.shape})")

        st.write("▶️ Step 3: Computing channel diffs…")
        t2: float = time.time()
        calibration_factors: List[float] = [
            0.002850975, 0.002861057, 0.002860953, 0.002837607, 
            0.00291866, 0.00295328, 0.00290534, 0.0029289
        ]
        df_new = loader.compute_channel_differences(df_new, calibration_factors)
        st.write(f"✅ compute_channel_differences done in {time.time() - t2:.2f}s")
        
        df_new.to_csv('Analysed crack data.csv', index = False)
        
        st.write("▶️ Step 4: Querying latest datetime from DB…")
        t3: float = time.time()
        latest_dt: Optional[pd.Timestamp] = db.get_latest_datetime()
        st.write(f"✅ get_latest_datetime returned {latest_dt!r} in {time.time() - t3:.2f}s")

        st.write("▶️ Step 5: Filtering new rows & saving to DB…")
        t4: float = time.time()
        if latest_dt is not None:
            # Ensure the datetime column is the correct type before comparing.
            df_new['datetime'] = pd.to_datetime(df_new['datetime'])
            df_new = df_new[df_new["datetime"] > latest_dt]
            st.write(f"   • {len(df_new)} rows newer than {latest_dt}")
        
        if df_new.empty:
            st.warning("No new rows to insert.")
        else:
            db.save_to_db(df_new)
            st.write(f"✅ save_to_db done in {time.time() - t4:.2f}s (inserted {len(df_new)} rows)")

        # Update the session state and rerun the app to show the new data.
        st.session_state["last_upload"] = fingerprint
        st.success("🎉 Upload sequence complete.")
        st.rerun()

else:
    last: Optional[str] = st.session_state.get("last_upload")
    if last:
        st.info(f"Last uploaded file: {last}")
    else:
        st.info("⚠️ No file uploaded yet. Please select one above.")